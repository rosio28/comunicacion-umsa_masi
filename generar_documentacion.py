#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script para generar documentación del proyecto en formato APA en Word
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

NEGRO = RGBColor(0, 0, 0)

def crear_documento_apa():
    """Crea un documento Word con formato APA"""
    
    doc = Document()
    
    # Configurar márgenes (APA requiere 1 pulgada en todos los lados)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    def agregar_texto_apa(texto, titulo=False, bold=False):
        """Agrega texto con formato APA (Arial 12, negro)"""
        p = doc.add_paragraph(texto)
        p.paragraph_format.line_spacing = 1.5
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.font.color.rgb = NEGRO
            if bold or titulo:
                run.font.bold = True
        return p

    # Título del documento
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo_run = titulo.add_run("Sistema de Comunicación Social UMSA")
    titulo_run.font.size = Pt(12)
    titulo_run.font.bold = True
    titulo_run.font.name = 'Arial'
    titulo_run.font.color.rgb = NEGRO

    # Subtítulo
    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitulo_run = subtitulo.add_run("Documentación Técnica del Proyecto")
    subtitulo_run.font.size = Pt(12)
    subtitulo_run.font.bold = True
    subtitulo_run.font.name = 'Arial'
    subtitulo_run.font.color.rgb = NEGRO

    # Autor y fecha
    autor = doc.add_paragraph()
    autor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    autor_run = autor.add_run(f"Documentación generada: {datetime.now().strftime('%d de %B de %Y')}")
    autor_run.font.size = Pt(12)
    autor_run.font.name = 'Arial'
    autor_run.font.color.rgb = NEGRO

    # Tabla de contenidos (simulada)
    doc.add_heading('Tabla de Contenidos', level=1)
    toc_heading = doc.paragraphs[-1]
    for run in toc_heading.runs:
        run.font.color.rgb = NEGRO

    toc_items = [
        '1. Introducción',
        '2. Descripción del Proyecto',
        '3. Arquitectura del Sistema',
        '4. Backend',
        '5. Frontend',
        '6. Base de Datos',
        '7. Infraestructura y Despliegue',
        '8. Despliegue con Docker',
        '9. Guía de Uso',
        '10. Configuración Avanzada',
        '11. Mantenimiento y Troubleshooting',
        '12. Conclusiones'
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Inches(0.25)
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.font.color.rgb = NEGRO

    # Agregar salto de página
    doc.add_page_break()
    
    # 1. INTRODUCCIÓN
    doc.add_heading('1. Introducción', level=1)
    heading = doc.paragraphs[-1]
    for run in heading.runs:
        run.font.color.rgb = NEGRO

    intro_text = "El Sistema de Comunicación Social UMSA es una plataforma web integral diseñada para la gestión centralizada de contenido comunicacional de la Universidad Mayor de San Andrés. Este sistema facilita la publicación y administración de noticias, eventos, convocatorias, información de docentes, galerías multimedia y trámites administrativos. La plataforma está desarrollada con tecnologías modernas y escalables, implementando una arquitectura de dos capas (frontend y backend) completamente containerizada con Docker."
    agregar_texto_apa(intro_text)

    # Agregar salto de página
    doc.add_page_break()

    # 2. DESCRIPCIÓN DEL PROYECTO
    doc.add_heading('2. Descripción del Proyecto', level=1)
    heading = doc.paragraphs[-1]
    for run in heading.runs:
        run.font.color.rgb = NEGRO

    desc_text = "El proyecto nace de la necesidad institucional de centralizar y modernizar los canales de comunicación de la UMSA. La solución proporciona un panel de administración intuitivo para gestionar múltiples módulos de contenido, control de acceso basado en roles, autenticación segura mediante JWT, almacenamiento de archivos en Cloudinary, y envío de notificaciones por correo electrónico."
    agregar_texto_apa(desc_text)

    doc.add_heading('2.1 Objetivos Principales', level=2)
    heading = doc.paragraphs[-1]
    for run in heading.runs:
        run.font.color.rgb = NEGRO

    objetivos = [
        "Centralizar la información comunicacional institucional en una única plataforma.",
        "Proporcionar herramientas de administración de contenido intuitivas y eficientes.",
        "Implementar seguridad robusta mediante autenticación JWT y control de roles.",
        "Facilitar la gestión de eventos, noticias, docentes y documentos.",
        "Optimizar la experiencia del usuario con una interfaz moderna y responsiva.",
        "Garantizar escalabilidad y mantenibilidad mediante containerización con Docker."
    ]
    for obj in objetivos:
        p = doc.add_paragraph(obj)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Inches(0.25)
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.font.color.rgb = NEGRO

    # Agregar salto de página
    doc.add_page_break()
    
    # 3. ARQUITECTURA DEL SISTEMA
    doc.add_heading('3. Arquitectura del Sistema', level=1)
    heading = doc.paragraphs[-1]
    for run in heading.runs:
        run.font.color.rgb = NEGRO

    arch_text = "El sistema implementa una arquitectura de tres capas bien definidas: presentación (frontend React), lógica de negocio (backend PHP), y persistencia de datos (PostgreSQL). Esta separación permite desarrollo independiente, mejor mantenibilidad y escalabilidad horizontal."
    agregar_texto_apa(arch_text)

    doc.add_heading('3.1 Capas de la Aplicación', level=2)
    heading = doc.paragraphs[-1]
    for run in heading.runs:
        run.font.color.rgb = NEGRO

    capas_data = [
        ("Capa de Presentación", "Frontend React 18 con Vite, Tailwind CSS y componentes reutilizables. Interfaz responsiva compatible con dispositivos móviles y de escritorio."),
        ("Capa de Lógica de Negocio", "API RESTful en PHP 8.2 con arquitectura basada en controladores. Autenticación con JWT, validación de datos y gestión de permisos."),
        ("Capa de Persistencia", "PostgreSQL 15 con esquema relacional normalizado. Almacenamiento de usuarios, contenido, eventos y metadata del sistema.")
    ]

    for capa, desc in capas_data:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        capa_run = p.add_run(f"{capa}: ")
        capa_run.font.bold = True
        capa_run.font.color.rgb = NEGRO
        capa_run.font.name = 'Arial'
        capa_run.font.size = Pt(12)

        desc_run = p.add_run(desc)
        desc_run.font.name = 'Arial'
        desc_run.font.size = Pt(12)
        desc_run.font.color.rgb = NEGRO

    doc.add_heading('3.2 Flujo de Comunicación', level=2)
    heading = doc.paragraphs[-1]
    for run in heading.runs:
        run.font.color.rgb = NEGRO

    flujo_text = "El cliente frontend realiza solicitudes HTTP/HTTPS al servidor backend mediante la librería Axios. El backend procesa la solicitud, valida la autenticación si es requerida, consulta la base de datos, y retorna una respuesta JSON. Los archivos se almacenan en Cloudinary para optimización y servicio de CDN."
    agregar_texto_apa(flujo_text)

    # Agregar salto de página
    doc.add_page_break()
    
    # 4. BACKEND
    doc.add_heading('4. Backend', level=1)
    heading = doc.paragraphs[-1]
    for run in heading.runs:
        run.font.color.rgb = NEGRO

    backend_text = "El backend es una API RESTful desarrollada en PHP 8.2 que proporciona todos los servicios necesarios para la operación del sistema. Implementa un router centralizado, controladores especializados por dominio, middleware de autenticación, y utilidades reutilizables para respuestas, JWT, carga de archivos y envío de correos."
    agregar_texto_apa(backend_text)

    doc.add_heading('4.1 Tecnologías y Dependencias', level=2)
    heading = doc.paragraphs[-1]
    for run in heading.runs:
        run.font.color.rgb = NEGRO

    tech_items = [
        "PHP 8.2: Lenguaje de programación del servidor",
        "PostgreSQL 15: Sistema gestor de base de datos relacional",
        "Firebase PHP-JWT: Generación y validación de tokens JWT",
        "PHPMailer: Servicio de envío de correos electrónicos",
        "Cloudinary: Almacenamiento y optimización de imágenes",
        "Composer: Gestor de dependencias de PHP",
        "Docker: Containerización de la aplicación"
    ]
    for item in tech_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Inches(0.25)
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.font.color.rgb = NEGRO

    doc.add_heading('4.2 Estructura de Directorios', level=2)
    heading = doc.paragraphs[-1]
    for run in heading.runs:
        run.font.color.rgb = NEGRO

    struct_text = """backend/
├── api/
│   ├── controllers/     # Controladores por dominio funcional
│   ├── middleware/      # Middlewares (Auth.php)
│   ├── config/          # Configuración de BD y aplicación
│   ├── utils/           # JWT, Mailer, Upload, Response
│   ├── uploads/         # Almacenamiento local de archivos
│   ├── index.php        # Router principal
│   └── router.php       # Lógica de enrutamiento
├── Dockerfile          # Especificación de imagen Docker
└── composer.json       # Dependencias PHP y configuración"""

    p = doc.add_paragraph(struct_text)
    p.paragraph_format.line_spacing = 1.15
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        run.font.color.rgb = NEGRO

    doc.add_heading('4.3 Controladores Principales', level=2)
    heading = doc.paragraphs[-1]
    for run in heading.runs:
        run.font.color.rgb = NEGRO

    controllers = [
        ("AuthController", "Gestiona login, logout, cambio de contraseña, recuperación de contraseña, generación de tokens JWT y validación de sesiones."),
        ("DocentesController", "Administración de información de docentes: CRUD, asignación de materias, horarios y datos académicos."),
        ("NoticiasController", "Publicación, edición y gestión de noticias. Soporte para categorización, destacados, vistas y búsqueda."),
        ("GaleriaController", "Gestión de galerías de imágenes, carga de fotos, organización por categorías."),
        ("AllControllers", "Controlador multiuso para gestión de eventos, convocatorias, grupos WhatsApp y otros módulos.")
    ]

    for controller, desc in controllers:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        name_run = p.add_run(f"{controller}: ")
        name_run.font.bold = True
        name_run.font.color.rgb = NEGRO
        name_run.font.name = 'Arial'
        name_run.font.size = Pt(12)

        desc_run = p.add_run(desc)
        desc_run.font.name = 'Arial'
        desc_run.font.size = Pt(12)
        desc_run.font.color.rgb = NEGRO

    doc.add_heading('4.4 Autenticación y Seguridad', level=2)
    heading = doc.paragraphs[-1]
    for run in heading.runs:
        run.font.color.rgb = NEGRO

    auth_text = "La autenticación se realiza mediante JWT (JSON Web Tokens) con expiración configurable. Las contraseñas se almacenan hasheadas con BCRYPT (costo 12). El middleware Auth.php valida cada solicitud protegida verificando la presencia y validez del token JWT en la cabecera Authorization."
    agregar_texto_apa(auth_text)

    doc.add_heading('4.5 Endpoints Principales', level=2)
    heading = doc.paragraphs[-1]
    for run in heading.runs:
        run.font.color.rgb = NEGRO

    endpoints = [
        ("POST /api/auth/login", "Autenticación con email y contraseña. Retorna token JWT y datos del usuario."),
        ("GET /api/auth/me", "Obtiene información del usuario autenticado."),
        ("POST /api/auth/cambiar-password", "Cambio de contraseña (requiere autenticación)."),
        ("POST /api/noticias", "Crear nueva noticia (requiere permisos de editor)."),
        ("GET /api/noticias", "Obtener listado de noticias publicadas."),
        ("GET /api/docentes", "Obtener información de docentes."),
        ("GET /api/galeria", "Obtener galerías de imágenes."),
        ("POST /api/upload", "Carga de archivos a Cloudinary.")
    ]

    for endpoint, desc in endpoints:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        ep_run = p.add_run(endpoint)
        ep_run.font.bold = True
        ep_run.font.name = 'Courier New'
        ep_run.font.size = Pt(10)
        ep_run.font.color.rgb = NEGRO

        desc_run = p.add_run(f": {desc}")
        desc_run.font.name = 'Arial'
        desc_run.font.size = Pt(12)
        desc_run.font.color.rgb = NEGRO

    # Agregar salto de página
    doc.add_page_break()
    
    # 5. FRONTEND
    doc.add_heading('5. Frontend', level=1)
    heading = doc.paragraphs[-1]
    for run in heading.runs:
        run.font.color.rgb = NEGRO

    frontend_text = "El frontend es una aplicación web desarrollada con React 18, Vite como bundler/servidor, y Tailwind CSS para estilos. Proporciona una interfaz de usuario moderna, responsiva y optimizada. Incluye panel de administración para gestión de contenido, autenticación mediante JWT, y renderizado de páginas públicas."
    agregar_texto_apa(frontend_text)

    doc.add_heading('5.1 Dependencias Principales', level=2)
    heading = doc.paragraphs[-1]
    for run in heading.runs:
        run.font.color.rgb = NEGRO

    deps = [
        "React 18.3.1: Librería de componentes de interfaz",
        "Vite 5.3.5: Bundler moderno y servidor de desarrollo",
        "React Router DOM 6.26.0: Enrutamiento de aplicación",
        "Axios 1.7.2: Cliente HTTP para comunicación con API",
        "Tailwind CSS 3.4.7: Framework de estilos CSS",
        "React Hook Form 7.52.1: Gestión eficiente de formularios",
        "TanStack React Query 5.51.0: Caché y sincronización de datos",
        "Lucide React 0.400.0: Librería de iconos SVG",
        "FullCalendar 6.1.14: Componente de calendario interactivo",
        "React Quill 2.0.0: Editor de texto enriquecido",
        "React Hot Toast 2.4.1: Notificaciones toast",
        "JWT Decode 4.0.0: Decodificación de tokens JWT"
    ]

    for dep in deps:
        p = doc.add_paragraph(dep)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Inches(0.25)
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.font.color.rgb = NEGRO

    doc.add_heading('5.2 Estructura de Carpetas', level=2)
    heading = doc.paragraphs[-1]
    for run in heading.runs:
        run.font.color.rgb = NEGRO

    frontend_struct = """frontend/
├── src/
│   ├── components/
│   │   ├── admin/       # Componentes del panel de administración
│   │   ├── layout/      # Componentes de diseño (Header, Sidebar, Footer)
│   │   ├── ui/          # Componentes genéricos reutilizables
│   │   └── ProtectedRoute.jsx  # Rutas protegidas por autenticación
│   ├── pages/
│   │   ├── HomePage.jsx         # Página de inicio pública
│   │   ├── PublicPages.jsx      # Páginas públicas (noticias, eventos)
│   │   └── admin/               # Páginas del panel de administración
│   ├── services/
│   │   ├── api.js               # Cliente Axios autenticado
│   │   ├── apiPublic.js         # Cliente para rutas públicas
│   │   └── services.js          # Funciones de negocio
│   ├── context/
│   │   └── AuthContext.jsx      # Estado global de autenticación
│   ├── utils/
│   │   ├── helpers.js           # Funciones de utilidad
│   │   └── patchDriveImages.js # Optimización de imágenes
│   ├── App.jsx          # Componente raíz y rutas principales
│   ├── main.jsx         # Punto de entrada
│   └── index.css        # Estilos globales
├── vite.config.js       # Configuración de Vite
├── tailwind.config.js   # Configuración de Tailwind CSS
├── postcss.config.cjs   # Procesador de CSS
├── nginx.conf           # Configuración de servidor web
├── package.json         # Dependencias y scripts
├── Dockerfile           # Especificación de imagen Docker
└── index.html           # Plantilla HTML"""

    p = doc.add_paragraph(frontend_struct)
    p.paragraph_format.line_spacing = 1.15
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        run.font.color.rgb = NEGRO

    doc.add_heading('5.3 Módulos Principales', level=2)
    heading = doc.paragraphs[-1]
    for run in heading.runs:
        run.font.color.rgb = NEGRO

    modulos = [
        ("Autenticación", "Login, logout, recuperación de contraseña. Almacenamiento de tokens en localStorage."),
        ("Panel de Administración", "Gestión de noticias, eventos, docentes, galerías, convocatorias y usuarios."),
        ("Área Pública", "Visualización de noticias, eventos, directorio de docentes, galerías."),
        ("Calendario", "Visualización interactiva de eventos académicos."),
        ("Editor de Contenido", "Editor de texto enriquecido (WYSIWYG) para crear y editar noticias.")
    ]

    for modulo, desc in modulos:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        mod_run = p.add_run(f"{modulo}: ")
        mod_run.font.bold = True
        mod_run.font.color.rgb = NEGRO
        mod_run.font.name = 'Arial'
        mod_run.font.size = Pt(12)

        desc_run = p.add_run(desc)
        desc_run.font.name = 'Arial'
        desc_run.font.size = Pt(12)
        desc_run.font.color.rgb = NEGRO

    # Agregar salto de página
    doc.add_page_break()
    
    # 6. BASE DE DATOS
    doc.add_heading('6. Base de Datos', level=1)
    heading = doc.paragraphs[-1]
    for run in heading.runs:
        run.font.color.rgb = NEGRO

    db_text = "El sistema utiliza PostgreSQL 15 como gestor de base de datos relacional. La base de datos contiene esquemas bien estructurados organizados en dominios funcionales: autenticación, contenido dinámico, personas y transparencia."
    agregar_texto_apa(db_text)

    doc.add_heading('6.1 Parámetros de Conexión', level=2)
    heading = doc.paragraphs[-1]
    for run in heading.runs:
        run.font.color.rgb = NEGRO

    db_info = [
        "Base de datos: comunicacion_umsa",
        "Usuario: ccs_user",
        "Contraseña: 123456",
        "Puerto: 5432",
        "Host: db (en Docker) / localhost (local)"
    ]

    for info in db_info:
        p = doc.add_paragraph(info)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Inches(0.25)
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.font.color.rgb = NEGRO

    doc.add_heading('6.2 Dominios Principales', level=2)
    heading = doc.paragraphs[-1]
    for run in heading.runs:
        run.font.color.rgb = NEGRO

    dominios = [
        ("Dominio 1: Autenticación", "Tabla usuarios (id, nombre, email, password_hash, rol, ci, apellidos, teléfono, avatar). Tabla sesiones_log para auditoría de acceso."),
        ("Dominio 2: Contenido", "Tablas noticias, eventos, convocatorias, categorias, grupos_whatsapp. Soporte para clasificación, búsqueda y filtrado."),
        ("Dominio 3: Personas", "Tabla docentes con información académica. Tabla materias vinculadas a docentes por gestión."),
        ("Dominio 4: Multimedia", "Tablas para galerías de imágenes y archivos. Integración con Cloudinary.")
    ]

    for dominio, desc in dominios:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        dom_run = p.add_run(f"{dominio}: ")
        dom_run.font.bold = True
        dom_run.font.color.rgb = NEGRO
        dom_run.font.name = 'Arial'
        dom_run.font.size = Pt(12)

        desc_run = p.add_run(desc)
        desc_run.font.name = 'Arial'
        desc_run.font.size = Pt(12)
        desc_run.font.color.rgb = NEGRO

    doc.add_heading('6.3 Archivos de Esquema SQL', level=2)
    heading = doc.paragraphs[-1]
    for run in heading.runs:
        run.font.color.rgb = NEGRO

    schema_items = [
        ("database/schema.sql", "Definición completa del esquema relacional, tablas, índices y restricciones."),
        ("database/seed_formal_data.sql", "Datos iniciales y categorías predeterminadas."),
        ("database/nuevos_modulos.sql", "Extensiones del esquema para nuevos módulos y funcionalidades."),
        ("backend/database.sql", "Respaldo adicional del esquema.")
    ]

    for archivo, desc in schema_items:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        for run in p.runs:
            run.font.color.rgb = NEGRO

        arch_run = p.add_run(f"{archivo}: ")
        arch_run.font.bold = True
        arch_run.font.color.rgb = NEGRO
        arch_run.font.name = 'Courier New'
        arch_run.font.size = Pt(10)

        desc_run = p.add_run(desc)
        desc_run.font.name = 'Arial'
        desc_run.font.size = Pt(12)
        desc_run.font.color.rgb = NEGRO

    doc.add_heading('6.4 Características de Seguridad', level=2)
    heading = doc.paragraphs[-1]
    for run in heading.runs:
        run.font.color.rgb = NEGRO

    security = [
        "Contraseñas hasheadas con BCRYPT (costo 12)",
        "Foreign keys para integridad referencial",
        "Índices optimizados para búsquedas frecuentes",
        "Campos de auditoría (creado_en, actualizado_en)",
        "Campos booleanos para control de estado (activo, publicado, etc.)"
    ]

    for sec in security:
        p = doc.add_paragraph(sec)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Inches(0.25)
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.font.color.rgb = NEGRO

    # Agregar salto de página
    doc.add_page_break()
    
    # 7. INFRAESTRUCTURA Y DESPLIEGUE
    doc.add_heading('7. Infraestructura y Despliegue', level=1)
    heading = doc.paragraphs[-1]
    for run in heading.runs:
        run.font.color.rgb = NEGRO

    infra_text = "El proyecto utiliza Docker y Docker Compose para la containerización y orquestación de servicios. Esta arquitectura garantiza un despliegue consistente y reproducible en cualquier entorno (desarrollo, prueba, producción)."
    agregar_texto_apa(infra_text)

    doc.add_heading('7.1 Servicios Docker', level=2)
    heading = doc.paragraphs[-1]
    for run in heading.runs:
        run.font.color.rgb = NEGRO

    services_text = "El archivo docker-compose.yml define tres servicios principales que se orquestan conjuntamente:"
    agregar_texto_apa(services_text)

    doc.add_heading('7.1.1 Servicio db (PostgreSQL)', level=3)
    heading = doc.paragraphs[-1]
    for run in heading.runs:
        run.font.color.rgb = NEGRO

    db_service = [
        "Imagen: postgres:15-alpine",
        "Puerto: 5432",
        "Volumen: db_data para persistencia de datos",
        "Healthcheck: Verificación de disponibilidad del servicio",
        "Variables de entorno: Credenciales de base de datos configurables"
    ]

    for item in db_service:
        p = doc.add_paragraph(item)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Inches(0.25)
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.font.color.rgb = NEGRO

    doc.add_heading('7.1.2 Servicio backend (API PHP)', level=3)
    heading = doc.paragraphs[-1]
    for run in heading.runs:
        run.font.color.rgb = NEGRO

    backend_service = [
        "Build: Construido desde ./backend/Dockerfile",
        "Puerto: 8000 (mapeado al puerto 80 del contenedor)",
        "Dependencias: Requiere que db esté saludable",
        "Volúmenes: ./backend/uploads para almacenamiento local",
        "Restart policy: unless-stopped"
    ]

    for item in backend_service:
        p = doc.add_paragraph(item)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Inches(0.25)
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.font.color.rgb = NEGRO

    doc.add_heading('7.1.3 Servicio frontend (Nginx + React)', level=3)
    heading = doc.paragraphs[-1]
    for run in heading.runs:
        run.font.color.rgb = NEGRO

    frontend_service = [
        "Build: Construcción de React con Vite, servido con Nginx",
        "Puerto: 80 (puerto HTTP estándar)",
        "nginx.conf: Configuración del servidor web",
        "Restart policy: unless-stopped"
    ]

    for item in frontend_service:
        p = doc.add_paragraph(item)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Inches(0.25)
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.font.color.rgb = NEGRO

    doc.add_heading('7.2 Volúmenes', level=2)
    heading = doc.paragraphs[-1]
    for run in heading.runs:
        run.font.color.rgb = NEGRO

    volumes_text = [
        ("db_data", "Volumen de Docker para persistencia de datos PostgreSQL"),
        ("./backend/uploads", "Directorio local para almacenamiento de archivos"),
        ("./database/schema.sql", "Montaje del esquema para inicialización automática de BD")
    ]

    for volume, desc in volumes_text:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        vol_run = p.add_run(f"{volume}: ")
        vol_run.font.bold = True
        vol_run.font.color.rgb = NEGRO
        vol_run.font.name = 'Courier New'
        vol_run.font.size = Pt(10)

        desc_run = p.add_run(desc)
        desc_run.font.name = 'Arial'
        desc_run.font.size = Pt(12)
        desc_run.font.color.rgb = NEGRO

    # Agregar salto de página
    doc.add_page_break()

    # 8. DESPLIEGUE CON DOCKER
    doc.add_heading('8. Despliegue con Docker', level=1)
    heading = doc.paragraphs[-1]
    for run in heading.runs:
        run.font.color.rgb = NEGRO

    deploy_text = "Docker Compose permite levantar toda la infraestructura del proyecto con un único comando, garantizando que todos los servicios estén correctamente configurados y comunicados."
    agregar_texto_apa(deploy_text)

    doc.add_heading('8.1 Requisitos Previos', level=2)
    heading = doc.paragraphs[-1]
    for run in heading.runs:
        run.font.color.rgb = NEGRO

    reqs = [
        "Docker Desktop instalado (incluye Docker y Docker Compose)",
        "Git para clonar el repositorio",
        "Mínimo 2GB de RAM disponible",
        "Puertos 80, 8000 y 5432 disponibles"
    ]

    for req in reqs:
        p = doc.add_paragraph(req)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Inches(0.25)
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.font.color.rgb = NEGRO

    doc.add_heading('8.2 Pasos de Instalación e Inicio', level=2)
    heading = doc.paragraphs[-1]
    for run in heading.runs:
        run.font.color.rgb = NEGRO

    steps = [
        "Clonar el repositorio del proyecto",
        "Navegar al directorio raíz: cd proyecto",
        "Crear archivo .env con las variables de configuración",
        "Construir las imágenes: docker-compose build",
        "Levantar los servicios: docker-compose up -d",
        "Verificar que todos los servicios estén ejecutándose: docker-compose ps",
        "Acceder a http://localhost para el frontend",
        "Acceder a http://localhost:8000 para el backend"
    ]

    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Inches(0.25)

        num_run = p.add_run(f"{i}. ")
        num_run.font.bold = True
        num_run.font.color.rgb = NEGRO
        num_run.font.name = 'Arial'
        num_run.font.size = Pt(12)

        step_run = p.add_run(step)
        step_run.font.name = 'Arial'
        step_run.font.size = Pt(12)
        step_run.font.color.rgb = NEGRO

    doc.add_heading('8.3 Comandos Útiles', level=2)
    heading = doc.paragraphs[-1]
    for run in heading.runs:
        run.font.color.rgb = NEGRO

    commands = [
        ("docker-compose up -d", "Inicia los servicios en background"),
        ("docker-compose down", "Detiene y elimina los contenedores"),
        ("docker-compose logs -f [servicio]", "Visualiza logs en tiempo real"),
        ("docker-compose restart [servicio]", "Reinicia un servicio específico"),
        ("docker-compose exec backend php -v", "Ejecuta comando dentro del contenedor"),
        ("docker-compose ps", "Lista estado de los contenedores")
    ]

    for cmd, desc in commands:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Inches(0.25)

        cmd_run = p.add_run(cmd)
        cmd_run.font.bold = True
        cmd_run.font.color.rgb = NEGRO
        cmd_run.font.name = 'Courier New'
        cmd_run.font.size = Pt(10)

        desc_run = p.add_run(f": {desc}")
        desc_run.font.name = 'Arial'
        desc_run.font.size = Pt(12)
        desc_run.font.color.rgb = NEGRO

    # Agregar salto de página
    doc.add_page_break()

    # 9. GUÍA DE USO
    doc.add_heading('9. Guía de Uso', level=1)
    heading = doc.paragraphs[-1]
    for run in heading.runs:
        run.font.color.rgb = NEGRO

    doc.add_heading('9.1 Acceso al Sistema', level=2)
    heading = doc.paragraphs[-1]
    for run in heading.runs:
        run.font.color.rgb = NEGRO

    acceso_text = "Una vez que los servicios están levantados en Docker, se accede a http://localhost. El sistema solicitará credenciales de login. Se debe proporcionar email y contraseña de un usuario con permisos administrativos."
    agregar_texto_apa(acceso_text)

    doc.add_heading('9.2 Roles y Permisos del Sistema', level=2)
    heading = doc.paragraphs[-1]
    for run in heading.runs:
        run.font.color.rgb = NEGRO

    roles = [
        ("superadmin", "Acceso total al sistema, gestión de usuarios, configuraciones y auditoría."),
        ("admin", "Acceso a gestión de contenido y usuarios con limitaciones."),
        ("editor", "Crear y editar noticias, eventos, convocatorias sin acceso a administración."),
        ("colaborador", "Visualizar y crear contenido. Acceso limitado a módulos específicos.")
    ]

    for rol, permisos in roles:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Inches(0.25)
        rol_run = p.add_run(f"{rol}: ")
        rol_run.font.bold = True
        rol_run.font.color.rgb = NEGRO
        rol_run.font.name = 'Arial'
        rol_run.font.size = Pt(12)

        perm_run = p.add_run(permisos)
        perm_run.font.name = 'Arial'
        perm_run.font.size = Pt(12)
        perm_run.font.color.rgb = NEGRO

    # Agregar salto de página
    doc.add_page_break()

    # 10. CONFIGURACIÓN AVANZADA
    doc.add_heading('10. Configuración Avanzada', level=1)
    heading = doc.paragraphs[-1]
    for run in heading.runs:
        run.font.color.rgb = NEGRO

    doc.add_heading('10.1 Archivo .env - Variables Completas', level=2)
    heading = doc.paragraphs[-1]
    for run in heading.runs:
        run.font.color.rgb = NEGRO

    env_vars_full = [
        "# PostgreSQL",
        "POSTGRES_DB=comunicacion_umsa",
        "POSTGRES_USER=ccs_user",
        "POSTGRES_PASSWORD=123456",
        "",
        "# JWT",
        "JWT_SECRET=tu_clave_secreta_de_32_caracteres_minimo",
        "JWT_EXPIRY=28800",
        "",
        "# URLs",
        "FRONTEND_URL=http://localhost",
        "BACKEND_URL=http://localhost:8000",
        "",
        "# Correo Electrónico",
        "MAIL_HOST=smtp.gmail.com",
        "MAIL_PORT=587",
        "MAIL_USER=tu_email@gmail.com",
        "MAIL_PASS=tu_contraseña_aplicacion",
        "MAIL_FROM_NAME=Comunicación Social UMSA",
        "MAIL_REPLY_TO=soporte@umsa.edu.bo",
        "ADMIN_EMAIL=admin@umsa.edu.bo"
    ]

    for var in env_vars_full:
        p = doc.add_paragraph(var)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.left_indent = Inches(0.25)
        for run in p.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = NEGRO

    doc.add_heading('10.2 Configuración de Cloudinary', level=2)
    heading = doc.paragraphs[-1]
    for run in heading.runs:
        run.font.color.rgb = NEGRO

    cloudinary_text = "Cloudinary se configura en backend/api/config/config.php con credenciales de API. Los archivos se transforman y optimizan automáticamente. El servicio proporciona CDN global para imágenes."
    agregar_texto_apa(cloudinary_text)

    doc.add_heading('10.3 Configuración de JWT', level=2)
    heading = doc.paragraphs[-1]
    for run in heading.runs:
        run.font.color.rgb = NEGRO

    jwt_config = [
        "JWT_SECRET debe ser una clave fuerte de mínimo 32 caracteres",
        "JWT_EXPIRY controla el tiempo de sesión (28800 segundos = 8 horas)",
        "Cambiar SECRET en producción invalida todos los tokens activos",
        "Los tokens se envían en la cabecera Authorization: Bearer {token}"
    ]

    for item in jwt_config:
        p = doc.add_paragraph(item)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Inches(0.25)
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.font.color.rgb = NEGRO

    # Agregar salto de página
    doc.add_page_break()

    # 11. MANTENIMIENTO Y TROUBLESHOOTING
    doc.add_heading('11. Mantenimiento y Troubleshooting', level=1)
    heading = doc.paragraphs[-1]
    for run in heading.runs:
        run.font.color.rgb = NEGRO

    doc.add_heading('11.1 Problemas Comunes y Soluciones', level=2)
    heading = doc.paragraphs[-1]
    for run in heading.runs:
        run.font.color.rgb = NEGRO

    problemas = [
        ("Backend no conecta a BD", "Ver logs: docker-compose logs db. Verificar que db esté saludable con: docker-compose ps"),
        ("Errores 401 (No autenticado)", "Verificar JWT_SECRET en .env. Confirmar token en Authorization header."),
        ("Carga de archivos falla", "Revisar credenciales de Cloudinary. Verificar permisos en ./backend/uploads"),
        ("Frontend no carga CSS", "Ejecutar: docker-compose restart frontend. Limpiar caché del navegador."),
        ("Puerto ya está en uso", "Cambiar puertos en docker-compose.yml o terminar procesos que usan 80, 8000, 5432")
    ]

    for problema, solucion in problemas:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Inches(0.25)
        prob_run = p.add_run(f"Problema: {problema}")
        prob_run.font.bold = True
        prob_run.font.color.rgb = NEGRO
        prob_run.font.name = 'Arial'
        prob_run.font.size = Pt(12)

        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Inches(0.5)
        sol_run = p.add_run(f"Solución: {solucion}")
        sol_run.font.name = 'Arial'
        sol_run.font.size = Pt(12)
        sol_run.font.color.rgb = NEGRO

    doc.add_heading('11.2 Mantenimiento Periódico', level=2)
    heading = doc.paragraphs[-1]
    for run in heading.runs:
        run.font.color.rgb = NEGRO

    tareas = [
        "Revisar logs regularmente: docker-compose logs -f backend",
        "Hacer backup de BD: docker-compose exec db pg_dump -U ccs_user comunicacion_umsa > backup.sql",
        "Actualizar dependencias: composer update (backend), npm update (frontend)",
        "Monitorear espacio en disco de volúmenes Docker",
        "Validar integridad de datos en base de datos",
        "Limpiar archivos temporales en ./backend/tmp",
        "Revisar logs de acceso y errores del servidor"
    ]

    for tarea in tareas:
        p = doc.add_paragraph(tarea)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Inches(0.25)
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.font.color.rgb = NEGRO

    # Agregar salto de página
    doc.add_page_break()

    # 12. CONCLUSIONES
    doc.add_heading('12. Conclusiones', level=1)
    heading = doc.paragraphs[-1]
    for run in heading.runs:
        run.font.color.rgb = NEGRO
        "El Sistema de Comunicación Social UMSA es una solución integral y moderna que centraliza la gestión de contenido comunicacional institucional.",
        "La arquitectura de tres capas (presentación, lógica de negocio, datos) basada en tecnologías contemporáneas garantiza escalabilidad, mantenibilidad y seguridad.",
        "La containerización con Docker permite despliegue consistente y reproducible en cualquier entorno (desarrollo, prueba, producción).",
        "La separación clara entre frontend (React) y backend (PHP) permite desarrollo independiente y facilita el mantenimiento futuro.",
        "La autenticación basada en JWT proporciona seguridad robusta y permite la integración con otros sistemas.",
        "PostgreSQL como base de datos relacional garantiza integridad de datos y permite consultas complejas.",
        "Recomendaciones para futuras mejoras: implementar pruebas unitarias e integración continua, monitoreo en producción, documentación de APIs con OpenAPI/Swagger, caché distribuida (Redis), y análisis de logs centralizado."
    ]

    for conclusion in conclusiones:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Inches(0.25)
        conc_run = p.add_run(f"• {conclusion}")
        conc_run.font.name = 'Arial'
        conc_run.font.size = Pt(12)
        conc_run.font.color.rgb = NEGRO
    
    # Guardar documento
    ruta_salida = 'Documentacion_Proyecto_CCS_UMSA.docx'
    doc.save(ruta_salida)
    
    print(f"✓ Documento creado exitosamente: {ruta_salida}")
    print(f"✓ Formato: APA")
    print(f"✓ Fuente: Arial 12pt")
    print(f"✓ Interlineado: 1.5")
    print(f"✓ Ubicación: {ruta_salida}")

if __name__ == '__main__':
    crear_documento_apa()
