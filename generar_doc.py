#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

NEGRO = RGBColor(0, 0, 0)

def crear_documento():
    doc = Document()
    
    # Márgenes APA (1 pulgada en todos lados)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # PORTADA
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo_run = titulo.add_run("Sistema de Comunicación Social UMSA")
    titulo_run.font.size = Pt(12)
    titulo_run.font.bold = True
    titulo_run.font.name = 'Arial'
    titulo_run.font.color.rgb = NEGRO
    
    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitulo_run = subtitulo.add_run("Documentación Técnica del Proyecto")
    subtitulo_run.font.size = Pt(12)
    subtitulo_run.font.bold = True
    subtitulo_run.font.name = 'Arial'
    subtitulo_run.font.color.rgb = NEGRO
    
    fecha_p = doc.add_paragraph()
    fecha_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fecha_run = fecha_p.add_run(f"Documentación generada: {datetime.now().strftime('%d de %B de %Y')}")
    fecha_run.font.size = Pt(12)
    fecha_run.font.name = 'Arial'
    fecha_run.font.color.rgb = NEGRO
    
    doc.add_page_break()
    
    # TABLA DE CONTENIDOS
    toc = doc.add_heading('Tabla de Contenidos', level=1)
    for run in toc.runs:
        run.font.color.rgb = NEGRO
    
    toc_items = [
        '1. Introducción', '2. Descripción del Proyecto', '3. Arquitectura del Sistema',
        '4. Backend', '5. Frontend', '6. Base de Datos', '7. Infraestructura',
        '8. Despliegue con Docker', '9. Guía de Uso', '10. Configuración Avanzada',
        '11. Mantenimiento y Troubleshooting', '12. Conclusiones'
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Inches(0.25)
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.font.color.rgb = NEGRO
    
    doc.add_page_break()
    
    # SECCIONES
    seccion_1(doc)
    seccion_2(doc)
    seccion_3(doc)
    seccion_4(doc)
    seccion_5(doc)
    seccion_6(doc)
    seccion_7(doc)
    seccion_8(doc)
    seccion_9(doc)
    seccion_10(doc)
    seccion_11(doc)
    seccion_12(doc)
    
    # GUARDAR
    ruta_salida = 'Documentacion_Proyecto_CCS_UMSA.docx'
    doc.save(ruta_salida)
    print(f"✓ Documento creado: {ruta_salida}")
    print(f"✓ Formato: APA (Arial 12pt, interlineado 1.5)")
    print(f"✓ Todo en negro, sin espacios innecesarios")

def agregar_encabezado(doc, titulo, nivel=1):
    h = doc.add_heading(titulo, level=nivel)
    for run in h.runs:
        run.font.color.rgb = NEGRO

def agregar_parrafo(doc, texto):
    p = doc.add_paragraph(texto)
    p.paragraph_format.line_spacing = 1.5
    for run in p.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.font.color.rgb = NEGRO

def agregar_lista(doc, items):
    for item in items:
        p = doc.add_paragraph(item)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Inches(0.25)
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.font.color.rgb = NEGRO

def seccion_1(doc):
    agregar_encabezado(doc, "1. Introducción", 1)
    agregar_parrafo(doc, "El Sistema de Comunicación Social UMSA es una plataforma web integral diseñada para la gestión centralizada de contenido comunicacional de la Universidad Mayor de San Andrés. Este sistema facilita la publicación y administración de noticias, eventos, convocatorias, información de docentes, galerías multimedia y trámites administrativos. La plataforma está desarrollada con tecnologías modernas y escalables, implementando una arquitectura de dos capas completamente containerizada con Docker.")
    doc.add_page_break()

def seccion_2(doc):
    agregar_encabezado(doc, "2. Descripción del Proyecto", 1)
    agregar_parrafo(doc, "El proyecto nace de la necesidad institucional de centralizar y modernizar los canales de comunicación de la UMSA. La solución proporciona un panel de administración intuitivo, control de acceso basado en roles, autenticación segura mediante JWT, almacenamiento en Cloudinary y envío de notificaciones por correo.")
    agregar_encabezado(doc, "2.1 Objetivos Principales", 2)
    objetivos = [
        "Centralizar la información comunicacional institucional en una única plataforma.",
        "Proporcionar herramientas de administración intuitivas y eficientes.",
        "Implementar seguridad robusta mediante JWT y control de roles.",
        "Facilitar la gestión de eventos, noticias, docentes y documentos.",
        "Optimizar la experiencia del usuario con interfaz moderna.",
        "Garantizar escalabilidad mediante containerización con Docker."
    ]
    agregar_lista(doc, objetivos)
    doc.add_page_break()

def seccion_3(doc):
    agregar_encabezado(doc, "3. Arquitectura del Sistema", 1)
    agregar_parrafo(doc, "El sistema implementa una arquitectura de tres capas: presentación (frontend React), lógica de negocio (backend PHP) y persistencia de datos (PostgreSQL). Esta separación permite desarrollo independiente, mejor mantenibilidad y escalabilidad.")
    agregar_encabezado(doc, "3.1 Capas de la Aplicación", 2)
    capas = [
        "Capa de Presentación: Frontend React 18 con Vite y Tailwind CSS.",
        "Capa de Lógica de Negocio: API RESTful en PHP 8.2 con JWT.",
        "Capa de Persistencia: PostgreSQL 15 con esquema relacional normalizado."
    ]
    agregar_lista(doc, capas)
    doc.add_page_break()

def seccion_4(doc):
    agregar_encabezado(doc, "4. Backend", 1)
    agregar_parrafo(doc, "El backend es una API RESTful en PHP 8.2 con router centralizado, controladores especializados, middleware de autenticación y utilidades reutilizables.")
    agregar_encabezado(doc, "4.1 Tecnologías", 2)
    techs = ["PHP 8.2", "PostgreSQL 15", "Firebase PHP-JWT", "PHPMailer", "Cloudinary", "Docker"]
    agregar_lista(doc, techs)
    agregar_encabezado(doc, "4.2 Controladores Principales", 2)
    controllers = [
        "AuthController: Autenticación, JWT y sesiones.",
        "DocentesController: Gestión de docentes y materias.",
        "NoticiasController: Publicación de noticias.",
        "GaleriaController: Gestión de imágenes.",
        "AllControllers: Eventos, convocatorias y otros."
    ]
    agregar_lista(doc, controllers)
    doc.add_page_break()

def seccion_5(doc):
    agregar_encabezado(doc, "5. Frontend", 1)
    agregar_parrafo(doc, "Aplicación web con React 18, Vite y Tailwind CSS. Proporciona interfaz moderna, responsiva y optimizada con panel de administración.")
    agregar_encabezado(doc, "5.1 Dependencias Principales", 2)
    deps = [
        "React 18.3.1", "Vite 5.3.5", "React Router DOM 6.26", "Axios 1.7.2",
        "Tailwind CSS 3.4.7", "React Hook Form 7.52", "React Query 5.51", "FullCalendar 6.1.14"
    ]
    agregar_lista(doc, deps)
    doc.add_page_break()

def seccion_6(doc):
    agregar_encabezado(doc, "6. Base de Datos", 1)
    agregar_parrafo(doc, "PostgreSQL 15 con esquema relacional normalizado organizado en dominios funcionales: autenticación, contenido, personas y multimedia.")
    agregar_encabezado(doc, "6.1 Parámetros de Conexión", 2)
    params = ["Base de datos: comunicacion_umsa", "Usuario: ccs_user", "Contraseña: 123456", "Puerto: 5432"]
    agregar_lista(doc, params)
    agregar_encabezado(doc, "6.2 Características de Seguridad", 2)
    seguridad = ["Contraseñas hasheadas con BCRYPT (costo 12)", "Foreign keys para integridad", "Índices optimizados", "Campos de auditoría", "Control de estado"]
    agregar_lista(doc, seguridad)
    doc.add_page_break()

def seccion_7(doc):
    agregar_encabezado(doc, "7. Infraestructura", 1)
    agregar_parrafo(doc, "Docker y Docker Compose para containerización. Garantiza despliegue consistente en cualquier entorno.")
    agregar_encabezado(doc, "7.1 Servicios Docker", 2)
    servicios = [
        "db: PostgreSQL 15 con volumen db_data.",
        "backend: API PHP en puerto 8000.",
        "frontend: React con Nginx en puerto 80."
    ]
    agregar_lista(doc, servicios)
    doc.add_page_break()

def seccion_8(doc):
    agregar_encabezado(doc, "8. Despliegue con Docker", 1)
    agregar_encabezado(doc, "8.1 Requisitos", 2)
    reqs = ["Docker Desktop instalado", "Git para clonar", "Mínimo 2GB RAM", "Puertos 80, 8000, 5432 disponibles"]
    agregar_lista(doc, reqs)
    agregar_encabezado(doc, "8.2 Pasos", 2)
    pasos = [
        "1. Clonar repositorio",
        "2. Crear archivo .env",
        "3. docker-compose build",
        "4. docker-compose up -d",
        "5. Verificar con docker-compose ps"
    ]
    agregar_lista(doc, pasos)
    doc.add_page_break()

def seccion_9(doc):
    agregar_encabezado(doc, "9. Guía de Uso", 1)
    agregar_parrafo(doc, "Acceder a http://localhost. El sistema solicita credenciales. Se debe proporcionar email y contraseña con permisos administrativos.")
    agregar_encabezado(doc, "9.1 Roles y Permisos", 2)
    roles = [
        "superadmin: Acceso total.",
        "admin: Gestión de contenido y usuarios.",
        "editor: Crear y editar contenido.",
        "colaborador: Visualizar y crear con limitaciones."
    ]
    agregar_lista(doc, roles)
    doc.add_page_break()

def seccion_10(doc):
    agregar_encabezado(doc, "10. Configuración Avanzada", 1)
    agregar_encabezado(doc, "10.1 Archivo .env", 2)
    env_text = "Variables necesarias: POSTGRES_DB, POSTGRES_USER, POSTGRES_PASSWORD, JWT_SECRET, JWT_EXPIRY, FRONTEND_URL, BACKEND_URL, MAIL_HOST, MAIL_PORT, MAIL_USER, MAIL_PASS, MAIL_FROM_NAME, ADMIN_EMAIL."
    agregar_parrafo(doc, env_text)
    agregar_encabezado(doc, "10.2 JWT", 2)
    agregar_parrafo(doc, "JWT_SECRET debe tener mínimo 32 caracteres. JWT_EXPIRY en segundos (28800 = 8 horas). Los tokens se envían en Authorization header.")
    doc.add_page_break()

def seccion_11(doc):
    agregar_encabezado(doc, "11. Mantenimiento y Troubleshooting", 1)
    agregar_encabezado(doc, "11.1 Problemas Comunes", 2)
    problemas = [
        "Backend no conecta a BD: Ver docker-compose logs db",
        "Errores 401: Verificar JWT_SECRET y Authorization header",
        "Carga de archivos falla: Revisar credenciales de Cloudinary",
        "Frontend no carga CSS: docker-compose restart frontend",
        "Puerto en uso: Cambiar puertos en docker-compose.yml"
    ]
    agregar_lista(doc, problemas)
    agregar_encabezado(doc, "11.2 Tareas Periódicas", 2)
    tareas = [
        "Revisar logs: docker-compose logs -f backend",
        "Backup BD: docker-compose exec db pg_dump",
        "Actualizar dependencias",
        "Monitorear espacio en disco",
        "Limpiar archivos temporales"
    ]
    agregar_lista(doc, tareas)
    doc.add_page_break()

def seccion_12(doc):
    agregar_encabezado(doc, "12. Conclusiones", 1)
    conclusiones = [
        "Solución integral y moderna para gestión de comunicación institucional.",
        "Arquitectura escalable garantiza mantenibilidad y seguridad.",
        "Docker permite despliegue consistente en cualquier entorno.",
        "Separación de capas facilita desarrollo independiente.",
        "JWT proporciona seguridad robusta.",
        "PostgreSQL garantiza integridad de datos.",
        "Futuras mejoras: pruebas unitarias, CI/CD, monitoreo, OpenAPI/Swagger."
    ]
    agregar_lista(doc, conclusiones)

if __name__ == '__main__':
    crear_documento()
